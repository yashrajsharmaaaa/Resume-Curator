"""
File processing utilities for Resume Curator application.

This module provides text extraction from various file formats
including PDF, DOC, and DOCX files. Designed for SDE1 portfolio
demonstration with clean error handling.
"""

import io
import logging
from typing import Optional

from fastapi import UploadFile
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded file.
    
    Supports PDF, DOC, and DOCX file formats.
    
    Args:
        file: Uploaded file to process
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If text extraction fails
    """
    try:
        # Read file content
        content = await file.read()
        file.file.seek(0)  # Reset file pointer
        
        # Determine file type from filename
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            return extract_text_from_pdf(content)
        elif filename.endswith('.docx'):
            return extract_text_from_docx(content)
        elif filename.endswith('.doc'):
            # For .doc files, we'll try to handle them as best we can
            # In a production environment, you might want to use python-docx2txt
            # or convert to .docx first
            return extract_text_from_doc(content)
        else:
            raise Exception(f"Unsupported file type: {filename}")
            
    except Exception as e:
        logger.error(f"Text extraction failed for {file.filename}: {e}")
        raise Exception(f"Could not extract text from {file.filename}: {str(e)}")


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF file content.
    
    Args:
        content: PDF file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If PDF processing fails
    """
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(stream=content, filetype="pdf")
        
        text_content = []
        
        # Extract text from each page
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if text.strip():
                text_content.append(text.strip())
        
        pdf_document.close()
        
        # Join all pages with double newlines
        full_text = '\n\n'.join(text_content)
        
        if not full_text.strip():
            raise Exception("No text content found in PDF")
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        raise Exception(f"PDF processing error: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX file content.
    
    Args:
        content: DOCX file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If DOCX processing fails
    """
    try:
        # Open DOCX from bytes
        doc_stream = io.BytesIO(content)
        document = Document(doc_stream)
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_content.append(text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        # Join all content with newlines
        full_text = '\n'.join(text_content)
        
        if not full_text.strip():
            raise Exception("No text content found in DOCX")
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        raise Exception(f"DOCX processing error: {str(e)}")


def extract_text_from_doc(content: bytes) -> str:
    """
    Extract text from DOC file content.
    
    Note: This is a simplified implementation. For production use,
    consider using more robust libraries like python-docx2txt or
    converting DOC to DOCX first.
    
    Args:
        content: DOC file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If DOC processing fails
    """
    try:
        # For now, we'll return an error message suggesting DOCX format
        # In a production environment, you would implement proper DOC parsing
        # or use a service to convert DOC to DOCX
        
        raise Exception(
            "DOC format support is limited. Please convert to DOCX format for better text extraction."
        )
        
    except Exception as e:
        logger.error(f"DOC text extraction failed: {e}")
        raise Exception(f"DOC processing error: {str(e)}")


def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text content
    """
    if not text:
        return ""
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive whitespace
    lines = []
    for line in text.split('\n'):
        cleaned_line = ' '.join(line.split())  # Normalize spaces
        if cleaned_line:  # Skip empty lines
            lines.append(cleaned_line)
    
    # Join lines with single newlines
    cleaned_text = '\n'.join(lines)
    
    # Remove excessive consecutive newlines
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text.strip()


def get_text_statistics(text: str) -> dict:
    """
    Get basic statistics about extracted text.
    
    Args:
        text: Text content to analyze
        
    Returns:
        Dictionary with text statistics
    """
    if not text:
        return {
            "character_count": 0,
            "word_count": 0,
            "line_count": 0,
            "paragraph_count": 0
        }
    
    # Basic counts
    character_count = len(text)
    word_count = len(text.split())
    line_count = len(text.split('\n'))
    paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
    
    return {
        "character_count": character_count,
        "word_count": word_count,
        "line_count": line_count,
        "paragraph_count": paragraph_count
    }